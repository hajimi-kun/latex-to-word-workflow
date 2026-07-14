from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

# English + common Chinese bibliography section titles
REFERENCE_HEADINGS = frozenset({
    "references",
    "bibliography",
    "works cited",
    "参考文献",
    "引用文献",
    "参考资料",
})


def ensure_style(doc: Document, name: str, base: str = "Normal"):
    if name in doc.styles:
        return doc.styles[name]
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles[base]
    return style


def image_only(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing | .//w:pict")) and not paragraph.text.strip()


def is_reference_heading(text: str) -> bool:
    normalized = re.sub(r"\s+", " ", text).strip().rstrip(":：").lower()
    return normalized in REFERENCE_HEADINGS or text.strip().rstrip(":：") in REFERENCE_HEADINGS


def main() -> None:
    parser = argparse.ArgumentParser(description="Map generated DOCX content to semantic Word styles.")
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    output = args.output or args.input

    doc = Document(args.input)
    figure = ensure_style(doc, "Figure")
    caption = ensure_style(doc, "Caption")
    table_body = ensure_style(doc, "Table Body")
    bibliography = ensure_style(doc, "Bibliography")

    figure.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.first_line_indent = Cm(0)
    figure.paragraph_format.keep_with_next = True
    bibliography.font.size = bibliography.font.size or Pt(9)
    if bibliography.paragraph_format.left_indent is None:
        bibliography.paragraph_format.left_indent = Cm(0.68)
        bibliography.paragraph_format.first_line_indent = Cm(-0.68)

    previous_was_image = False
    in_references = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if is_reference_heading(text):
            in_references = True
            previous_was_image = False
            continue
        if image_only(paragraph):
            paragraph.style = figure
            previous_was_image = True
        elif previous_was_image and text:
            paragraph.style = caption
            previous_was_image = False
        elif in_references and re.match(r"^(\[\d+\]|\(\d+\))\s", text):
            paragraph.style = bibliography
        elif text:
            previous_was_image = False

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = table_body

    for section in doc.sections:
        usable_width = section.page_width - section.left_margin - section.right_margin
        for shape in doc.inline_shapes:
            if shape.width and shape.width > usable_width:
                ratio = shape.height / shape.width
                shape.width = usable_width
                shape.height = int(usable_width * ratio)

    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()

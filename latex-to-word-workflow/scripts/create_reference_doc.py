from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def field_run(paragraph, instruction: str, display: str, *, bold: bool = False):
    begin = paragraph.add_run()
    if bold:
        begin.bold = True
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin._r.append(fld_begin)

    code = paragraph.add_run()
    if bold:
        code.bold = True
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    code._r.append(instr)

    separate = paragraph.add_run()
    if bold:
        separate.bold = True
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    separate._r.append(fld_separate)

    result = paragraph.add_run(display)
    result.bold = bold
    no_proof = OxmlElement("w:noProof")
    result._r.get_or_add_rPr().append(no_proof)

    end = paragraph.add_run()
    if bold:
        end.bold = True
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end._r.append(fld_end)


def bookmarked_field(paragraph, bookmark: str, bookmark_id: int,
                     instruction: str, display: str, *, bold: bool = False):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark)
    paragraph._p.append(start)
    field_run(paragraph, instruction, display, bold=bold)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def add_math(paragraph):
    math = OxmlElement("m:oMath")
    for text in ("E", "=", "m"):
        run = OxmlElement("m:r")
        token = OxmlElement("m:t")
        token.text = text
        run.append(token)
        math.append(run)
    superscript = OxmlElement("m:sSup")
    base = OxmlElement("m:e")
    base_run = OxmlElement("m:r")
    base_token = OxmlElement("m:t")
    base_token.text = "c"
    base_run.append(base_token)
    base.append(base_run)
    exponent = OxmlElement("m:sup")
    exponent_run = OxmlElement("m:r")
    exponent_token = OxmlElement("m:t")
    exponent_token.text = "2"
    exponent_run.append(exponent_token)
    exponent.append(exponent_run)
    superscript.extend((base, exponent))
    math.append(superscript)
    paragraph._p.append(math)


def style(doc: Document, name: str, size: float, *, bold: bool = False,
          alignment=None, first_line: float = 0, before: float = 0, after: float = 0):
    item = doc.styles[name] if name in doc.styles else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    item.font.name = "Times New Roman"
    item.font.size = Pt(size)
    item.font.bold = bold
    item.paragraph_format.alignment = alignment
    item.paragraph_format.first_line_indent = Cm(first_line)
    item.paragraph_format.space_before = Pt(before)
    item.paragraph_format.space_after = Pt(after)
    item.quick_style = True
    return item


def main() -> None:
    parser = argparse.ArgumentParser(description="Create the generic reference.docx bundled with this skill.")
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    args.output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(2.54)

    normal = style(doc, "Normal", 10.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    body = style(doc, "Body Text", 10.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=0.5, after=3)
    body.base_style = normal
    style(doc, "Title", 15, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    style(doc, "Subtitle", 10.5, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    for name, size, before in (("Heading 1", 13, 14), ("Heading 2", 11.5, 10), ("Heading 3", 10.5, 8)):
        heading = style(doc, name, size, bold=True, before=before, after=4)
        heading.paragraph_format.keep_with_next = True
    figure = style(doc, "Figure", 10.5, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=3)
    figure.paragraph_format.keep_with_next = True
    caption = style(doc, "Caption", 9, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    style(doc, "Equation", 10.5, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
    style(doc, "Table Body", 9)
    bibliography = style(doc, "Bibliography", 9)
    bibliography.paragraph_format.left_indent = Cm(0.68)
    bibliography.paragraph_format.first_line_indent = Cm(-0.68)

    samples = [
        ("Title", "Example Manuscript Title"),
        ("Subtitle", "Author names and affiliations"),
        ("Heading 1", "1. Introduction"),
        ("Heading 2", "1.1. Background"),
        ("Heading 3", "1.1.1. Detailed heading"),
        ("Body Text", "This paragraph demonstrates the manuscript body style, including first-line indentation, justification, and paragraph spacing."),
    ]
    for name, text in samples:
        doc.add_paragraph(text, style=name)

    crossref = doc.add_paragraph(style="Body Text")
    crossref.add_run("Native Word cross-reference examples: Fig. ")
    field_run(crossref, "REF RefFig001 \\h", "1")
    crossref.add_run(", Table ")
    field_run(crossref, "REF RefTable001 \\h", "1")
    crossref.add_run(", and Eq. (")
    field_run(crossref, "REF RefEquation001 \\h", "1")
    crossref.add_run(").")

    equation = doc.add_paragraph(style="Equation")
    add_math(equation)
    equation.add_run("    (")
    bookmarked_field(equation, "RefEquation001", 1, "SEQ Equation \\* ARABIC", "1")
    equation.add_run(")")

    figure_path = Path(__file__).resolve().parents[1] / "examples" / "minimal" / "figure.png"
    figure_paragraph = doc.add_paragraph(style="Figure")
    figure_paragraph.add_run().add_picture(str(figure_path), width=Cm(11.5))
    figure_caption = doc.add_paragraph(style="Caption")
    figure_caption.add_run("Fig. ").bold = True
    bookmarked_field(figure_caption, "RefFig001", 2, "SEQ Fig. \\* ARABIC", "1", bold=True)
    figure_caption.add_run(". ").bold = True
    figure_caption.add_run("Example figure caption demonstrating native numbering and caption style.")

    table_caption = doc.add_paragraph(style="Caption")
    table_caption.add_run("Table ").bold = True
    bookmarked_field(table_caption, "RefTable001", 3, "SEQ Table \\* ARABIC", "1", bold=True)
    table_caption.add_run(". ").bold = True
    table_caption.add_run("Example table caption demonstrating native numbering and caption style.")

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.text = ("Variable", "Condition", "Value")[column_index] if row_index == 0 else ("Example", "Baseline", "1.0")[column_index]
            for paragraph in cell.paragraphs:
                paragraph.style = "Table Body"
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "E7EAF0")
                cell._tc.get_or_add_tcPr().append(shading)
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph(
        "[1] A. Author, B. Author, Example article title, Journal 10 (2025) 100-110. https://doi.org/10.0000/example.",
        style="Bibliography",
    )

    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from lxml import etree


REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
INPUT_RE = re.compile(r"\\(?:input|include)\s*\{([^}]+)\}")
GRAPHICS_RE = re.compile(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}")
# BBT / zotero.lua often embeds citationKey or citation-key in the field JSON
ZOTERO_CITEKEY_RE = re.compile(
    r'"(?:citationKey|citation-key|citekey)"\s*:\s*"([^"]+)"',
    re.I,
)


def strip_comments(text: str) -> str:
    cleaned: list[str] = []
    for line in text.splitlines():
        match = re.search(r"(?<!\\)%", line)
        cleaned.append(line[: match.start()] if match else line)
    return "\n".join(cleaned)


def resolve_include(parent: Path, raw_name: str) -> Path:
    candidate = (parent / raw_name.strip()).resolve()
    return candidate if candidate.suffix else candidate.with_suffix(".tex")


def expand_tex(path: Path, stack: tuple[Path, ...] = ()) -> str:
    resolved = path.resolve()
    if resolved in stack:
        chain = " -> ".join(str(item) for item in (*stack, resolved))
        raise RuntimeError(f"Cyclic LaTeX include: {chain}")
    if not resolved.exists():
        raise FileNotFoundError(f"LaTeX source not found: {resolved}")
    source = strip_comments(resolved.read_text(encoding="utf-8"))
    parts: list[str] = []
    cursor = 0
    for match in INPUT_RE.finditer(source):
        parts.append(source[cursor : match.start()])
        parts.append(expand_tex(resolve_include(resolved.parent, match.group(1)), (*stack, resolved)))
        cursor = match.end()
    parts.append(source[cursor:])
    return "".join(parts)


def tex_counts(entry: Path) -> dict:
    text = expand_tex(entry)
    graphics = GRAPHICS_RE.findall(text)
    return {
        "figure_environments": len(re.findall(r"\\begin\{figure\*?\}", text)),
        "table_environments": len(re.findall(r"\\begin\{table\*?\}", text)),
        "includegraphics": len(graphics),
        "pdf_includegraphics": [g for g in graphics if Path(g).suffix.lower() == ".pdf"],
        "extensionless_includegraphics": [g for g in graphics if not Path(g).suffix],
    }


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    parser = argparse.ArgumentParser(
        description="Validate DOCX structure, relationships, content counts, and Zotero fields."
    )
    parser.add_argument("docx", type=Path)
    parser.add_argument("--expect-zotero", type=int, help="Expected live Zotero citation field count")
    parser.add_argument(
        "--expect-zotero-keys",
        nargs="+",
        metavar="KEY",
        help="Citation keys that must appear inside Zotero field JSON (Better BibTeX citekeys)",
    )
    parser.add_argument("--expect-tables", type=int, help="Expected Word table count")
    parser.add_argument("--expect-images", type=int, help="Expected inline image count")
    parser.add_argument(
        "--tex",
        type=Path,
        help="LaTeX entry point: compare source figure/table/includegraphics counts to the DOCX",
    )
    parser.add_argument(
        "--no-enforce-tex-counts",
        action="store_true",
        help="With --tex, only report source counts (do not fail on mismatch)",
    )
    parser.add_argument(
        "--allow-pdf-figures",
        action="store_true",
        help="With --tex, do not fail when \\includegraphics uses an explicit .pdf path",
    )
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()

    report: dict = {"path": str(args.docx.resolve()), "zip_valid": False, "missing_relationships": []}
    try:
        with ZipFile(args.docx) as archive:
            report["zip_valid"] = True
            xml = etree.fromstring(archive.read("word/document.xml"))
            rels = etree.fromstring(archive.read("word/_rels/document.xml.rels"))
            used = set(xml.xpath("//@r:id | //@r:embed | //@r:link", namespaces={"r": REL_NS}))
            defined = {node.get("Id") for node in rels}
            report["missing_relationships"] = sorted(used - defined)
            document_text = archive.read("word/document.xml").decode("utf-8", errors="replace")
            report["zotero_citation_fields"] = document_text.count("ADDIN ZOTERO_ITEM CSL_CITATION")
            report["zotero_citation_keys"] = sorted(set(ZOTERO_CITEKEY_RE.findall(document_text)))
            report["unresolved_pandoc_cite_markers"] = document_text.count("[@")
    except (BadZipFile, KeyError, etree.XMLSyntaxError) as exc:
        report["zip_error"] = str(exc)

    try:
        doc = Document(args.docx)
        report.update({
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "images": len(doc.inline_shapes),
            "styles": sorted({paragraph.style.name for paragraph in doc.paragraphs if paragraph.style}),
        })
    except Exception as exc:
        report["python_docx_error"] = str(exc)

    failures: list[str] = []
    if args.tex:
        try:
            counts = tex_counts(args.tex)
        except (OSError, RuntimeError) as exc:
            report["tex_error"] = str(exc)
            failures.append(f"could not expand LaTeX entry point: {exc}")
            counts = None
        if counts is not None:
            report["tex_counts"] = {
                "figure_environments": counts["figure_environments"],
                "table_environments": counts["table_environments"],
                "includegraphics": counts["includegraphics"],
            }
            report["pdf_includegraphics"] = counts["pdf_includegraphics"]
            report["extensionless_includegraphics"] = counts["extensionless_includegraphics"]
            if counts["pdf_includegraphics"]:
                report["pdf_figure_warning"] = (
                    "Word cannot embed PDF figures; change \\includegraphics to PNG/JPEG "
                    "(a same-stem .png next to figure.pdf is not used automatically)"
                )
            if counts["extensionless_includegraphics"]:
                report["extensionless_figure_warning"] = (
                    "Extensionless \\includegraphics paths may resolve to PDF; prefer explicit .png/.jpg"
                )
            if not args.no_enforce_tex_counts:
                docx_tables = report.get("tables")
                docx_images = report.get("images")
                if docx_tables is not None and docx_tables != counts["table_environments"]:
                    failures.append(
                        f"table count mismatch: LaTeX table envs={counts['table_environments']}, "
                        f"DOCX tables={docx_tables}"
                    )
                if docx_images is not None and docx_images != counts["includegraphics"]:
                    failures.append(
                        f"image count mismatch: LaTeX includegraphics={counts['includegraphics']}, "
                        f"DOCX images={docx_images}"
                    )
            if not args.allow_pdf_figures and counts["pdf_includegraphics"]:
                failures.append("PDF \\includegraphics paths present (not Word-safe)")

    if not report["zip_valid"]:
        failures.append("invalid ZIP/DOCX package")
    if report["missing_relationships"]:
        failures.append("missing OpenXML relationships")
    if "python_docx_error" in report:
        failures.append("python-docx could not load the document")
    if args.expect_zotero is not None and report.get("zotero_citation_fields") != args.expect_zotero:
        failures.append(f"expected {args.expect_zotero} Zotero fields")
    if args.expect_zotero_keys:
        found_keys = set(report.get("zotero_citation_keys") or [])
        missing_keys = [key for key in args.expect_zotero_keys if key not in found_keys]
        if missing_keys:
            failures.append(f"missing Zotero citation keys in fields: {missing_keys}")
    if args.expect_tables is not None and report.get("tables") != args.expect_tables:
        failures.append(f"expected {args.expect_tables} tables, found {report.get('tables')}")
    if args.expect_images is not None and report.get("images") != args.expect_images:
        failures.append(f"expected {args.expect_images} images, found {report.get('images')}")
    if report.get("unresolved_pandoc_cite_markers"):
        failures.append("unresolved pandoc cite markers [@...] remain in DOCX")
    report["ok"] = not failures
    report["failures"] = failures
    rendered = json.dumps(report, ensure_ascii=False, indent=2)
    if args.output:
        args.output.write_text(rendered, encoding="utf-8")
    print(rendered)
    raise SystemExit(0 if report["ok"] else 1)


if __name__ == "__main__":
    main()
